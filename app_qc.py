import streamlit as st
from docx import Document
import re
import PyPDF2
import io
import os

# -----------------------------
# 🔹 Normalisation
# -----------------------------
def normalize(text):
    return text.strip().lower()

def normalize_key(key):
    """
    Normalise la clé pour comparaison :
    - tout en minuscules
    - enlève caractères spéciaux au début et à la fin
    """
    key = key.strip().lower()
    key = re.sub(r'^[^a-z0-9]+', '', key)
    key = re.sub(r'[^a-z0-9]+$', '', key)
    return key

# -----------------------------
# 🔹 Parser boutons radio
# -----------------------------
def parse_radio(line):
    matches = re.findall(r'([◉●○])\s*(Yes|No)', line)
    for symbol, value in matches:
        if symbol in ["◉", "●"]:
            return value
    return None

# -----------------------------
# 🔹 Parser PDF
# -----------------------------
def parse_pdf(file):
    """
    Extrait le texte d'un fichier PDF
    """
    try:
        # Lire le fichier PDF
        pdf_reader = PyPDF2.PdfReader(file)
        text = ""
        for page in pdf_reader.pages:
            extracted_text = page.extract_text()
            if extracted_text:
                text += extracted_text
        
        # Parser le texte ligne par ligne comme pour DOCX
        data = {}
        for line in text.split('\n'):
            line = line.strip()
            if ":" in line:
                parts = line.split(":", 1)
                if len(parts) == 2:
                    key, rest = parts
                    key = normalize(key)
                    if "Yes" in rest and "No" in rest:
                        value = parse_radio(rest)
                    else:
                        value = rest.strip()
                    if value:
                        data[key] = normalize(value)
        
        return data
    except Exception as e:
        st.error(f"Erreur lors de la lecture du PDF: {str(e)}")
        return {}

def parse_pdf_tables(file):
    """
    Extrait le texte d'un PDF contenant des tables
    Note: L'extraction de tables depuis PDF est limitée
    """
    try:
        pdf_reader = PyPDF2.PdfReader(file)
        all_text = ""
        for page in pdf_reader.pages:
            extracted_text = page.extract_text()
            if extracted_text:
                all_text += extracted_text
        
        # Simuler la structure de table pour la compatibilité
        datasets = []
        
        # Détection basique de lignes avec séparateurs
        lines = all_text.split('\n')
        for line in lines:
            if '|' in line or '\t' in line:
                # Tenter de parser comme une ligne de tableau
                parts = re.split(r'[|\t]', line)
                if len(parts) > 1:
                    row_dict = {}
                    for i, part in enumerate(parts):
                        row_dict[f"colonne_{i}"] = normalize(part.strip())
                    if row_dict:
                        normalized_row = {normalize_key(k): v for k, v in row_dict.items()}
                        datasets.append(normalized_row)
        
        return datasets if datasets else [{}]  # Retourne un dict vide si pas de tables détectées
    except Exception as e:
        st.error(f"Erreur lors de la lecture du PDF: {str(e)}")
        return [{}]

# -----------------------------
# 🔹 Parser document 1 (DOCX ou PDF)
# -----------------------------
@st.cache_data
def parse_doc1(file, file_type):
    try:
        if file_type == "docx":
            doc = Document(file)
            data = {}
            for para in doc.paragraphs:
                text = para.text.strip()
                if ":" in text:
                    parts = text.split(":", 1)
                    if len(parts) == 2:
                        key, rest = parts
                        key = normalize(key)
                        if "Yes" in rest and "No" in rest:
                            value = parse_radio(rest)
                        else:
                            value = rest.strip()
                        if value:
                            data[key] = normalize(value)
            return data
        elif file_type == "pdf":
            return parse_pdf(file)
        else:
            return {}
    except Exception as e:
        st.error(f"Erreur lors du parsing du document 1: {str(e)}")
        return {}

# -----------------------------
# 🔹 Parser document 2 (DOCX ou PDF)
# -----------------------------
@st.cache_data
def parse_doc2(file, file_type):
    try:
        if file_type == "docx":
            doc = Document(file)
            datasets = []
            for table in doc.tables:
                if len(table.rows) > 0:
                    headers = [normalize(cell.text) for cell in table.rows[0].cells]
                    for row in table.rows[1:]:
                        values = [normalize(cell.text) for cell in row.cells]
                        if len(values) == len(headers):
                            row_dict = dict(zip(headers, values))
                            # Pré-normaliser les clés pour la recherche rapide
                            normalized_row = {normalize_key(k): v for k, v in row_dict.items()}
                            datasets.append(normalized_row)
            return datasets
        elif file_type == "pdf":
            return parse_pdf_tables(file)
        else:
            return []
    except Exception as e:
        st.error(f"Erreur lors du parsing du document 2: {str(e)}")
        return []

# -----------------------------
# 🔹 Comparaison optimisée
# -----------------------------
def compare_with_all_rows(doc1_data, datasets):
    best_match = None
    best_score = -1
    best_result = None

    normalized_doc1_keys = {k: normalize_key(k) for k in doc1_data.keys()}

    for row in datasets:
        score = 0
        result = {}

        for key, value in doc1_data.items():
            norm_key = normalized_doc1_keys[key]

            # Cherche une clé correspondante dans row
            matched_key = next((k for k in row.keys() if norm_key in k or k in norm_key), None)

            if not matched_key:
                result[key] = "❌ Colonne absente"
            elif row[matched_key] != value:
                result[key] = f"❌ Différent (doc1={value}, doc2={row[matched_key]})"
            else:
                result[key] = "✅ OK"
                score += 1

        if score > best_score:
            best_score = score
            best_match = row
            best_result = result

    return best_match, best_result, best_score

# -----------------------------
# 🔹 Streamlit UI
# -----------------------------
st.set_page_config(page_title="Document Comparator", layout="wide")

# Custom CSS for better styling
st.markdown("""
    <style>
    .stAlert {
        margin-top: 1rem;
        margin-bottom: 1rem;
    }
    </style>
    """, unsafe_allow_html=True)

st.title("📄 Comparateur de Documents Optimisé")
st.write("Upload les deux documents à comparer (Word ou PDF)")

col1, col2 = st.columns(2)

with col1:
    doc1_file = st.file_uploader("Document 1 (texte + radios)", type=["docx", "pdf"], key="doc1")
    doc1_type = None
    if doc1_file:
        doc1_type = "docx" if doc1_file.name.endswith('.docx') else "pdf"

with col2:
    doc2_file = st.file_uploader("Document 2 (tables)", type=["docx", "pdf"], key="doc2")
    doc2_type = None
    if doc2_file:
        doc2_type = "docx" if doc2_file.name.endswith('.docx') else "pdf"

if doc1_file and doc2_file and doc1_type and doc2_type:
    with st.spinner("Traitement des documents en cours..."):
        doc1_data = parse_doc1(doc1_file, doc1_type)
        datasets = parse_doc2(doc2_file, doc2_type)
        
        if doc1_data and datasets:
            best_match, result, score = compare_with_all_rows(doc1_data, datasets)

            st.subheader("📊 Résumé")
            st.write(f"Score de correspondance : **{score} / {len(doc1_data)}**")
            
            # Progress bar for visual feedback
            progress = score / len(doc1_data) if len(doc1_data) > 0 else 0
            st.progress(progress)

            st.subheader("🔍 Détails des comparaisons")
            
            # Create expandable sections for better organization
            success_items = []
            error_items = []
            
            for key, value in result.items():
                if "✅" in value:
                    success_items.append((key, value))
                else:
                    error_items.append((key, value))
            
            if success_items:
                with st.expander("✅ Éléments correspondants", expanded=False):
                    for key, value in success_items:
                        st.success(f"{key} → {value}")
            
            if error_items:
                with st.expander("❌ Éléments non correspondants", expanded=True):
                    for key, value in error_items:
                        st.error(f"{key} → {value}")

            st.subheader("📋 Ligne correspondante dans le document 2")
            if best_match:
                st.json(best_match)
            else:
                st.warning("Aucune correspondance trouvée")
        else:
            if not doc1_data:
                st.error("❌ Aucune donnée extraite du document 1. Vérifiez le format du fichier.")
            if not datasets or (len(datasets) == 1 and not datasets[0]):
                st.error("❌ Aucune donnée extraite du document 2. Vérifiez le format du fichier.")

elif doc1_file or doc2_file:
    if not doc1_file:
        st.info("📁 Veuillez uploader le document 1")
    if not doc2_file:
        st.info("📁 Veuillez uploader le document 2")

# Footer
st.markdown("---")
st.markdown("💡 **Conseil :** Pour de meilleurs résultats, assurez-vous que vos documents sont bien formatés avec des séparateurs (:) pour le document 1 et des tableaux pour le document 2.")